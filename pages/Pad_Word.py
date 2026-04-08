import streamlit as st
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import re
import io
import copy

# ---------------------------------------------------------------------------
# Padrões regex para identificação hierárquica de títulos
# ---------------------------------------------------------------------------
# Aceita numeração com ou sem espaço após o ponto: "1. TEXTO" ou "1.TEXTO"
REGEX_ITEM_PRINCIPAL  = re.compile(r'^\s*\d+\.\s*[A-Za-zÀ-ÖØ-öø-ÿ]')
REGEX_SUBITEM         = re.compile(r'^\s*\d+\.\d+\s+[A-Za-zÀ-ÖØ-öø-ÿ]')
REGEX_SUB_SUBITEM     = re.compile(r'^\s*\d+\.\d+\.\d+\s+[A-Za-zÀ-ÖØ-öø-ÿ]')
REGEX_BULLET_GAMBIARRA = re.compile(r'^\s*[-•]\s+')
# Bullet nível 2 - "o item": exige que o texto seja curto (< 60 chars) para não pegar início de frase
REGEX_BULLET_NIVEL2   = re.compile(r'^\s*[oO]\s+\S')
REGEX_LETRA_GAMBIARRA = re.compile(r'^\s*[a-z]\)\s+')

# Títulos de seção fixos que devem ser tratados como item principal mesmo sem numeração
TITULOS_FIXOS = {
    "OBJETIVO", "ABRANGÊNCIA", "ABRANGENCIA",
    "DOCUMENTO REFERÊNCIA", "DOCUMENTO REFERENCIA",
    "DEFINIÇÕES", "DEFINICOES",
    "RESPONSABILIDADES E ATRIBUIÇÕES", "RESPONSABILIDADES E ATRIBUICOES",
    "DESCRIÇÃO DO PROCESSO", "DESCRICAO DO PROCESSO",
    "ANEXOS", "ANEXO",
}

# Estilos do Word que representam passos com letra (a), b), c)...)
ESTILOS_LETRA = {"*PARÁGRAFO", "*PARAGRAFO", "Título P (IT)", "Titulo P (IT)"}


# ---------------------------------------------------------------------------
# Helpers de manipulação de runs
# ---------------------------------------------------------------------------

def aplicar_fonte_runs(paragrafo, negrito=False, tamanho_pt=12, nome_fonte='Arial',
                        italico=False, sublinhado=False):
    """
    Aplica formatação de fonte em todos os runs de um parágrafo SEM apagar o texto.
    Reseta italic e underline salvo se explicitamente solicitados.
    """
    for run in paragrafo.runs:
        run.font.name    = nome_fonte
        run.font.size    = Pt(tamanho_pt)
        run.font.bold    = negrito
        run.font.italic  = italico
        run.font.underline = sublinhado


def transformar_texto_runs(paragrafo, transformacao):
    """
    Aplica uma função de transformação de string (upper/lower/capitalize) em
    cada run individualmente, preservando a quantidade de runs e sua formatação.
    """
    for run in paragrafo.runs:
        run.text = transformacao(run.text)


def primeira_letra_maiuscula_runs(paragrafo):
    """
    Garante que a primeira letra do conteúdo do parágrafo seja maiúscula,
    operando diretamente nos runs sem apagar nada.
    """
    primeiro_char_encontrado = False
    for run in paragrafo.runs:
        if not run.text or primeiro_char_encontrado:
            # Passa pelo texto em minúsculo se ainda não achou a primeira letra
            if not primeiro_char_encontrado:
                run.text = run.text.lower()
            continue
        texto_lower = run.text.lower()
        for i, ch in enumerate(texto_lower):
            if ch.isalpha():
                run.text = texto_lower[:i] + texto_lower[i].upper() + texto_lower[i+1:]
                primeiro_char_encontrado = True
                break
        else:
            run.text = texto_lower  # run sem letra, só converte para lower


# ---------------------------------------------------------------------------
# Helper para adicionar borda ao parágrafo (usada em imagens)
# ---------------------------------------------------------------------------

def adicionar_borda_paragrafo(paragrafo, cor_hex="000000", tamanho=4, espaco=1):
    """
    Adiciona borda preta fina ao redor do parágrafo via XML (pBdr).
    Necessário para imagens conforme regra da IT.
    """
    pPr = paragrafo._p.get_or_add_pPr()

    # Remove pBdr anterior se existir para não duplicar
    for old_pBdr in pPr.findall(qn('w:pBdr')):
        pPr.remove(old_pBdr)

    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    for lado in ('top', 'left', 'bottom', 'right'):
        borda = etree.SubElement(pBdr, qn(f'w:{lado}'))
        borda.set(qn('w:val'),   'single')
        borda.set(qn('w:sz'),    str(tamanho))   # largura em 1/8 pt
        borda.set(qn('w:space'), str(espaco))
        borda.set(qn('w:color'), cor_hex)


def remover_borda_paragrafo(paragrafo):
    """Remove qualquer pBdr existente no parágrafo."""
    pPr = paragrafo._p.find(qn('w:pPr'))
    if pPr is not None:
        for pBdr in pPr.findall(qn('w:pBdr')):
            pPr.remove(pBdr)


# ---------------------------------------------------------------------------
# Função principal de classificação do parágrafo
# ---------------------------------------------------------------------------

def classificar_paragrafo(paragrafo):
    """
    Retorna uma string indicando a categoria do parágrafo:
      'sub_subitem'  → 6.1.1 (maiúsculo, sem negrito)
      'subitem'      → 6.1   (minúsculo com 1ª maiúscula, negrito)
      'item_principal' → 1.  (maiúsculo, negrito)
      'letra'        → a) b) c) (corpo, alinhado à esquerda, recuo 0,64)
      'lista_n1'     → bullet nível 1 (recuo 0,64)
      'lista_n2'     → bullet nível 2 / "o" (recuo 1,91)
      'imagem'       → parágrafo contém drawing
      'corpo'        → texto normal
    """
    texto  = paragrafo.text.strip()
    estilo = paragrafo.style.name
    xml    = paragrafo._p.xml

    tem_imagem = 'w:drawing' in xml

    # --- Imagem (parágrafo dedicado, texto praticamente vazio) ---
    if tem_imagem and len(texto) < 5:
        return 'imagem'

    # --- Hierarquia por texto (tem precedência sobre estilo para títulos) ---
    if REGEX_SUB_SUBITEM.match(texto):
        return 'sub_subitem'

    if REGEX_SUBITEM.match(texto):
        return 'subitem'

    # Heading 2 sem numeração visível também é subitem
    if estilo in ('Heading 2',) and texto and not REGEX_ITEM_PRINCIPAL.match(texto):
        return 'subitem'

    if REGEX_ITEM_PRINCIPAL.match(texto):
        return 'item_principal'

    # Título fixo sem numeração (ex: "OBJETIVO" em Heading 1)
    if texto.upper().strip() in TITULOS_FIXOS or estilo == 'Heading 1':
        if texto:
            return 'item_principal'

    # --- Passos com letra: estilo customizado OU regex ---
    if estilo in ESTILOS_LETRA or REGEX_LETRA_GAMBIARRA.match(texto):
        return 'letra'

    # --- Listas ---
    if estilo in ('List Paragraph',) or estilo.startswith('List'):
        # Tenta detectar nível via numPr
        try:
            nivel = paragrafo._p.pPr.numPr.ilvl.val
            return 'lista_n2' if nivel >= 1 else 'lista_n1'
        except AttributeError:
            return 'lista_n1'

    if REGEX_BULLET_NIVEL2.match(texto) and len(texto) < 60:
        return 'lista_n2'

    if REGEX_BULLET_GAMBIARRA.match(texto):
        return 'lista_n1'

    return 'corpo'


# ---------------------------------------------------------------------------
# Aplicador de regras por categoria
# ---------------------------------------------------------------------------

def formatar_paragrafo(paragrafo):
    """
    Aplica as regras de formatação da IT conforme a categoria do parágrafo.
    Nunca usa p.text = ... para não apagar runs existentes.
    """
    texto     = paragrafo.text.strip()
    categoria = classificar_paragrafo(paragrafo)

    # --- Espaçamento e recuo padrão (aplicado a todos) ---
    fmt = paragrafo.paragraph_format
    fmt.space_before        = Pt(0)
    fmt.space_after         = Pt(0)
    fmt.line_spacing        = 1.5
    fmt.line_spacing_rule   = WD_LINE_SPACING.MULTIPLE
    fmt.left_indent         = Cm(0)
    fmt.right_indent        = Cm(0)
    fmt.first_line_indent   = Cm(0)
    # Reseta paginação por padrão (sobrescrito nas categorias que precisam)
    fmt.keep_with_next      = False
    fmt.keep_together       = False

    # --- Aplicação por categoria ---

    if categoria == 'sub_subitem':
        # 6.1.1 → MAIÚSCULO, sem negrito, alinhado à esquerda
        # keep_with_next: título não fica solto no fim da página
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.keep_with_next  = True
        transformar_texto_runs(paragrafo, str.upper)
        aplicar_fonte_runs(paragrafo, negrito=False)

    elif categoria == 'subitem':
        # 5.1 → Primeira letra maiúscula, resto minúsculo, negrito, esquerda
        # keep_with_next: título não fica solto no fim da página
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.keep_with_next  = True
        primeira_letra_maiuscula_runs(paragrafo)
        aplicar_fonte_runs(paragrafo, negrito=True)

    elif categoria == 'item_principal':
        # 1. OBJETIVO → MAIÚSCULO, negrito, alinhado à esquerda
        # keep_with_next: título não fica solto no fim da página
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.keep_with_next  = True
        transformar_texto_runs(paragrafo, str.upper)
        aplicar_fonte_runs(paragrafo, negrito=True)

    elif categoria == 'letra':
        # a) b) c) → corpo normal, esquerda, recuo 0,64
        # keep_with_next: o texto do passo não fica separado da imagem que o ilustra
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent     = Cm(0.64)
        fmt.keep_with_next  = True
        aplicar_fonte_runs(paragrafo, negrito=False)

    elif categoria == 'lista_n1':
        # Bullet nível 1 → esquerda, recuo 0,64
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent     = Cm(0.64)
        aplicar_fonte_runs(paragrafo, negrito=False)

    elif categoria == 'lista_n2':
        # Bullet nível 2 (bolinha branca / "o") → esquerda, recuo 1,91
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent     = Cm(1.91)
        aplicar_fonte_runs(paragrafo, negrito=False)

    elif categoria == 'imagem':
        # Centralizado + borda preta fina (regra da IT)
        # keep_together: impede a imagem de ser cortada entre páginas
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.keep_together   = True
        adicionar_borda_paragrafo(paragrafo)

    else:
        # Corpo de texto → justificado, recuo zero
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        remover_borda_paragrafo(paragrafo)
        aplicar_fonte_runs(paragrafo, negrito=False)


# ---------------------------------------------------------------------------
# Formatação de tabelas
# ---------------------------------------------------------------------------

def formatar_tabelas(doc):
    """
    Centraliza tabelas e aplica fonte Arial 12 em todas as células.
    A primeira linha (cabeçalho) recebe negrito.
    """
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for idx_row, row in enumerate(table.rows):
            is_header = (idx_row == 0)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.font.bold = is_header


# ---------------------------------------------------------------------------
# Paginação: page break após a tabela de Elaboração/Aprovação
# ---------------------------------------------------------------------------

def _e_tabela_elaboracao(table):
    for row in table.rows:
        for cell in row.cells:
            txt = cell.text.strip().upper()
            if txt in ('ELABORAÇÃO', 'ELABORACAO', 'APROVAÇÃO', 'APROVACAO'):
                return True
    return False


def inserir_pagebreak_apos_tabela_elaboracao(doc):
    from docx.table import Table as DocxTable
    body = doc.element.body
    elementos = list(body)

    for i, el in enumerate(elementos):
        if el.tag.split('}')[-1] != 'tbl':
            continue
        if not _e_tabela_elaboracao(DocxTable(el, doc)):
            continue

        # Coletar todos os parágrafos logo após a tabela até o primeiro com conteúdo
        paragrafos_pos_tabela = []
        idx_primeiro_conteudo = None

        for j in range(i + 1, len(elementos)):
            proximo = elementos[j]
            if proximo.tag.split('}')[-1] != 'p':
                continue
            texto = ''.join(t.text or '' for t in proximo.iter(qn('w:t'))).strip()
            paragrafos_pos_tabela.append(proximo)
            if texto and idx_primeiro_conteudo is None:
                idx_primeiro_conteudo = len(paragrafos_pos_tabela) - 1
                break

        if not paragrafos_pos_tabela:
            break

        # Remover os parágrafos vazios que precedem o de conteúdo
        # (eles causam linhas em branco no topo da nova página)
        for p_vazio in paragrafos_pos_tabela[:idx_primeiro_conteudo]:
            txt = ''.join(t.text or '' for t in p_vazio.iter(qn('w:t'))).strip()
            tem_img = 'w:drawing' in p_vazio.xml if hasattr(p_vazio, 'xml') else False
            if not txt and not tem_img:
                body.remove(p_vazio)

        # Colocar pageBreakBefore no parágrafo com conteúdo
        alvo = paragrafos_pos_tabela[idx_primeiro_conteudo] if idx_primeiro_conteudo is not None else paragrafos_pos_tabela[0]

        pPr = alvo.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            alvo.insert(0, pPr)

        for old in pPr.findall(qn('w:pageBreakBefore')):
            pPr.remove(old)

        pbr = OxmlElement('w:pageBreakBefore')
        pbr.set(qn('w:val'), '1')
        pPr.append(pbr)
        break


# ---------------------------------------------------------------------------
# Paginação: zerar space_before em parágrafos que abrem página
# ---------------------------------------------------------------------------

def _tem_pagebreak_before(p_el):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is not None:
        pbr = pPr.find(qn('w:pageBreakBefore'))
        if pbr is not None and pbr.get(qn('w:val'), '1') not in ('0', 'false'):
            return True
    for br in p_el.iter(qn('w:br')):
        if br.get(qn('w:type')) == 'page':
            return True
    return False


def remover_space_before_inicio_pagina(doc):
    for el in doc.element.body.iter(qn('w:p')):
        if not _tem_pagebreak_before(el):
            continue
        pPr = el.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            el.insert(0, pPr)
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:beforeAutospacing'), '0')


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

def padronizar_documento(arquivo_upload):
    doc = Document(arquivo_upload)

    for p in doc.paragraphs:
        texto = p.text.strip()
        tem_imagem = 'w:drawing' in p._p.xml

        # Ignora parágrafos completamente vazios sem imagem
        if not texto and not tem_imagem:
            continue

        formatar_paragrafo(p)

    formatar_tabelas(doc)

    # Garante page break após tabela de Elaboração/Aprovação
    inserir_pagebreak_apos_tabela_elaboracao(doc)

    # Remove espaçamento no topo de páginas
    remover_space_before_inicio_pagina(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Interface Streamlit
# ---------------------------------------------------------------------------

st.subheader("📄 Padronizador de ITs — QSMS")
st.write("Faça o upload do arquivo bruto. O sistema aplicará automaticamente as regras de formatação da empresa.")

arquivo_upload = st.file_uploader(
    "Suba o arquivo .docx aqui",
    type=["docx"],
    key="upload_doc_qsms"
)

if arquivo_upload:
    if st.button("Processar Documento", type="primary"):
        try:
            with st.spinner("Padronizando documento..."):
                buffer = padronizar_documento(arquivo_upload)

            st.success("Documento padronizado com sucesso.")
            st.download_button(
                label="📥 Baixar Documento Padronizado",
                data=buffer.getvalue(),
                file_name=f"Padronizado_{arquivo_upload.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"Falha na execução: {e}")
            st.exception(e)